"""Generates docs/solid_audit.docx from the audit data."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
RED    = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xE2, 0x6B, 0x0A)
YELLOW = RGBColor(0xBF, 0x8F, 0x00)
GREEN  = RGBColor(0x37, 0x86, 0x10)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1F, 0x37, 0x63)
LIGHT_BLUE  = RGBColor(0xDD, 0xE8, 0xF0)
HEADER_BLUE = RGBColor(0x1F, 0x37, 0x63)

SEVERITY_COLOURS = {
    "Major":    RED,
    "Moderate": ORANGE,
    "Minor":    YELLOW,
    "Clean":    GREEN,
}

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, colour: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    hex_col = f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}"
    shd.set(qn("w:fill"), hex_col)
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def add_header_row(table, headers, bg: RGBColor = HEADER_BLUE):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)


def add_module_table(doc, module_name, rows):
    """rows: list of (principle, finding, severity_str)"""
    doc.add_heading(module_name, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    add_header_row(table, ["Principle", "Finding", "Severity"])
    set_col_width(table, 0, 2.2)
    set_col_width(table, 1, 11.0)
    set_col_width(table, 2, 2.3)

    for i, (principle, finding, severity) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = principle
        cells[1].text = finding
        cells[2].text = severity

        for cell in cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.text else None

        colour = SEVERITY_COLOURS.get(severity, YELLOW)
        set_cell_bg(cells[2], colour)
        run = cells[2].paragraphs[0].runs
        if run:
            run[0].font.color.rgb = WHITE
            run[0].bold = True
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i % 2 == 0:
            set_cell_bg(cells[0], LIGHT_BLUE)
            set_cell_bg(cells[1], LIGHT_BLUE)

    doc.add_paragraph()


def add_cross_cutting_table(doc, title, rows, col_headers):
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1 + len(rows), cols=len(col_headers))
    table.style = "Table Grid"
    add_header_row(table, col_headers)

    widths = {2: [4.0, 11.5], 3: [4.0, 7.5, 4.0], 4: [4.5, 4.5, 3.5, 3.0]}
    for i, w in enumerate(widths.get(len(col_headers), [])):
        set_col_width(table, i, w)

    for ri, row_data in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            if cells[ci].paragraphs[0].runs:
                cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if ri % 2 == 0:
            for cell in cells:
                set_cell_bg(cell, LIGHT_BLUE)

    doc.add_paragraph()


# ── Document ───────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title = doc.add_heading("SOLID Audit — Aegis AI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Date: June 7, 2026").italic = True
    doc.add_paragraph(
        "Scope: All source modules — every file read in full before this audit was written."
    )

    # ── MODULE-BY-MODULE ──────────────────────────────────────────────────────
    doc.add_heading("Module-by-Module Findings", level=1)

    modules = [
        ("app/config.py", [
            ("S", "Single responsibility: reads env vars, exposes constants.", "Clean"),
            ("O", "Adding a new config value is additive — no modification needed.", "Clean"),
            ("D", "os.getenv() is the correct abstraction boundary here.", "Clean"),
            ("I", "No violations.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("app/llm.py", [
            ("S", "Two concerns: LLM factory (build_llm) and message builder (build_messages) — related but distinct.", "Minor"),
            ("O", "Only supports OpenAI — adding another provider means modifying the function.", "Minor"),
            ("D", "build_llm() returns the concrete ChatOpenAI type, not BaseChatModel — callers are tied to the concrete type.", "Minor"),
            ("I", "No violations.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("app/rag_chain.py", [
            ("S", "build_rag_response() does: retrieval, message construction, LLM creation, streaming, and fallback — all in one function.", "Moderate"),
            ("O", "hybrid flag exposes retrieval strategy to callers.", "Minor"),
            ("D", "Creates ChatOpenAI inline (lines 41–47) instead of calling build_llm() from app/llm.py — the existing abstraction is ignored.", "Major"),
            ("I", "7 parameters — callers who don't use hybrid or selected_docs still see them.", "Minor"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("rag/retriever.py", [
            ("S", "build_retriever() handles both vector-only and hybrid retrieval — two strategies in one function.", "Minor"),
            ("O", "Adding a third retrieval strategy requires modifying build_retriever().", "Minor"),
            ("D", "OpenAIEmbeddings instantiated directly — embedding provider hardcoded.", "Minor"),
            ("I", "hybrid=False forces callers who only want vector retrieval to be aware of the BM25 option.", "Minor"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("rag/ingestion.py", [
            ("S", "9 distinct responsibilities in one file: OCR noise detection, table summarisation, page rendering, chart/vision summarisation, title-based chunking, basic PDF loading, enhanced PDF loading, ChromaDB storage, document listing.", "Major"),
            ("O", "ingest_file(enhanced=True/False) mixes two pipelines — adding a third strategy means modifying the function.", "Moderate"),
            ("D", "_build_llm() is a LOCAL copy of the LLM factory, separate from build_llm() in app/llm.py. OpenAIEmbeddings and Chroma also instantiated directly.", "Major"),
            ("I", "ingest_file(enhanced=...) serves two different caller types through one function.", "Minor"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("tools/tools.py", [
            ("S", "make_tools() defines two tools: search_knowledge_base and calculate_breach_cost — different dependencies, different concerns.", "Minor"),
            ("O", "Adding a new tool requires modifying make_tools() and its return list.", "Minor"),
            ("D", "build_retriever() called directly — acceptable coupling at this level.", "Clean"),
            ("I", "Each tool has a focused, minimal interface.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("tools/chain.py", [
            ("S", "_run_async() does: local tool setup, MCP setup, tool aggregation, LLM construction, tool binding, message history building, agentic loop (5 rounds), tool execution, tool call logging, content extraction.", "Major"),
            ("O", "MCP server configuration hardcoded in _run_async() — adding a new MCP server means modifying the function.", "Minor"),
            ("D", "ChatOpenAI instantiated inline. _CVE_SERVER path computed locally instead of from config.", "Major"),
            ("I", "run_with_tools() has 5 required parameters — minor.", "Minor"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("guardrails/prompt_injection.py", [
            ("S", "Single responsibility: score a message for injection risk.", "Clean"),
            ("O", "Threshold 0.7 is a module-level constant — appropriate.", "Clean"),
            ("D", "ChatOpenAI instantiated inline instead of using build_llm().", "Moderate"),
            ("I", "check_prompt_injection() takes a string, returns (bool, float) — focused.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("observability/logging_setup.py", [
            ("S", "Two concerns: (1) logging configuration (basicConfig, getLogger) and (2) the log_llm_call decorator.", "Minor"),
            ("O", "log_llm_call only supports timing — adding cost tracking requires modifying the function.", "Minor"),
            ("D", "No external dependencies beyond stdlib.", "Clean"),
            ("I", "No violations.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("observability/fault_tolerance.py", [
            ("S", "Two closely related primitives: retry and timeout — appropriate together.", "Clean"),
            ("O", "with_retry() hardcodes retry_if_exception_type(Exception) — very minor.", "Minor"),
            ("D", "Only stdlib + tenacity.", "Clean"),
            ("I", "No violations.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/breach_triage_agent.py", [
            ("S", "_run_agent_async() does: MCP setup, tool aggregation, LLM construction, agent creation, invocation, content extraction, tool call log building.", "Moderate"),
            ("O", "MCP config hardcoded.", "Minor"),
            ("D", "ChatOpenAI instantiated inline. _CVE_SERVER path computed locally. Content extraction duplicated inline.", "Major"),
            ("I", "run_agent() takes 4 params — clean.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/breach_workflow.py", [
            ("S", "Defines 3 agents (Assessment, Research, Report) + graph assembly + MCP setup in one file.", "Moderate"),
            ("O", "Adding a new workflow step requires modifying graph assembly.", "Minor"),
            ("D", "_make_llm() is a local copy of the LLM factory. _extract_text() is duplicated from other files. _CVE_SERVER computed locally.", "Major"),
            ("I", "Clean public interface.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/composed_workflow.py", [
            ("S", "Defines 6 agents/nodes + conditional routing + graph assembly in one file. critical_response_node and standard_response_node are ~50 lines of near-identical code.", "Major"),
            ("O", "severity_router() hardcodes 'critical' and 'high' as strings — adding a new severity path requires modifying the function.", "Moderate"),
            ("D", "_make_llm() local copy (different defaults than other files). _extract_text() duplicated. _CVE_SERVER computed locally.", "Major"),
            ("I", "Clean public interface.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/supervisor_workflow.py", [
            ("S", "Defines: _make_llm(), _extract_text(), _run_in_thread(), _make_specialist_tools() with 3 sub-agents as closures, supervisor prompt, _run_supervisor_async(), run_supervisor() — all in one file.", "Major"),
            ("O", "Adding a new specialist requires modifying _make_specialist_tools().", "Minor"),
            ("D", "_make_llm() local copy. ChatOpenAI also used inline in _run_supervisor_async(). _CVE_SERVER computed locally.", "Major"),
            ("I", "Clean public interface.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/supervisor_hitl.py", [
            ("S", "Duplicates ALL of supervisor_workflow.py's responsibilities PLUS adds _extract_severity(), requires_approval(), and two-phase public functions. The entire _make_specialist_tools() is copy-pasted from supervisor_workflow.py.", "Major"),
            ("O", "requires_approval() hardcodes ('critical', 'high') as the non-approval set. _extract_severity() hardcodes the recognised severity levels.", "Moderate"),
            ("D", "_make_llm() local copy (identical to supervisor_workflow.py). _extract_text() and _run_in_thread() also duplicated. _CVE_SERVER computed locally.", "Major"),
            ("I", "run_hitl_phase1() and run_hitl_phase2() have clean, focused interfaces.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/multimodal_agent.py", [
            ("S", "Single responsibility: enrich an incident with image analysis.", "Clean"),
            ("O", "No violations.", "Clean"),
            ("D", "ChatOpenAI instantiated inline with fixed temperature=0.0, max_tokens=1024.", "Moderate"),
            ("I", "enrich_with_image(incident, image_bytes, mime_type) — focused.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("agents/a2a_client.py", [
            ("S", "Three focused HTTP client functions — all related to one remote server.", "Clean"),
            ("O", "No violations.", "Clean"),
            ("D", "httpx is appropriate. A2A_SERVER_URL from env var is correct.", "Clean"),
            ("I", "Three focused, independent functions.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("mcp_servers/cve_server.py", [
            ("S", "Single responsibility: look up a CVE from the NVD API and format the result.", "Clean"),
            ("O", "CVSS version priority (v3.1 → v3.0 → v2) is a hardcoded loop — adding a new CVSS version requires modifying the function.", "Minor"),
            ("D", "NVD_API URL is a module-level constant — acceptable for a stable public API. httpx is appropriate.", "Minor"),
            ("I", "lookup_cve(cve_id: str) -> str — single parameter, single return. Focused.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("a2a_server/threat_intel_server.py", [
            ("S", "The run() route handler mixes HTTP concerns (request parsing, response model) with business logic (LLM construction, invocation, content extraction).", "Moderate"),
            ("O", "AGENT_CARD has 'url': 'http://localhost:8888' hardcoded — must change for any deployment.", "Minor"),
            ("D", "ChatOpenAI instantiated inline on every request. Content extraction duplicated inline — same pattern as _extract_text() across other modules.", "Major"),
            ("I", "Two minimal, focused endpoint schemas.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
        ("ui/app.py", [
            ("S", "460 lines. All 10 modes, sidebar rendering, HITL approval card, chat rendering, and tool call display are handled in one file. The largest SRP violation in the codebase by line count.", "Major"),
            ("O", "Adding a new mode requires modifying the if/elif chain in BOTH the sidebar section AND the chat handler section independently.", "Moderate"),
            ("D", "Imports and calls all agent modules directly — tightly coupled to concrete implementations. Acceptable at the UI boundary but worth noting.", "Minor"),
            ("I", "No violations.", "Clean"),
            ("L", "N/A — no class hierarchy.", "Clean"),
        ]),
    ]

    for mod_name, rows in modules:
        add_module_table(doc, mod_name, rows)

    # ── CROSS-CUTTING ─────────────────────────────────────────────────────────
    doc.add_heading("Cross-Cutting Violations", level=1)
    doc.add_paragraph(
        "These violations span multiple modules. Fixing them once fixes every affected module simultaneously."
    )

    add_cross_cutting_table(doc,
        "1. _make_llm() duplicated in 4 files — Major DIP + DRY",
        [
            ("agents/breach_workflow.py",    "0.0", "512"),
            ("agents/composed_workflow.py",  "0.0", "1024"),
            ("agents/supervisor_workflow.py","0.0", "1024"),
            ("agents/supervisor_hitl.py",    "0.0", "1024"),
        ],
        ["File", "Default temperature", "Default max_tokens"],
    )
    doc.add_paragraph(
        "app/llm.py already provides build_llm() — the intended factory. "
        "All four files ignore it and define their own copy with different default parameters. "
        "Fix: all modules call build_llm() from app/llm.py."
    )

    add_cross_cutting_table(doc,
        "2. _extract_text() duplicated in 5+ files — Major DRY",
        [
            ("agents/breach_workflow.py",          "Named _extract_text()"),
            ("agents/composed_workflow.py",         "Named _extract_text()"),
            ("agents/supervisor_workflow.py",       "Named _extract_text()"),
            ("agents/supervisor_hitl.py",           "Named _extract_text()"),
            ("agents/breach_triage_agent.py",       "Inline, not named"),
            ("a2a_server/threat_intel_server.py",   "Inline, not named"),
        ],
        ["File", "How it appears"],
    )
    doc.add_paragraph(
        "Fix: move to agents/utils.py as extract_text(content) -> str."
    )

    add_cross_cutting_table(doc,
        "3. Tool call log extraction duplicated in 5 files — Major DRY",
        [
            ("agents/breach_triage_agent.py",   "In _run_agent_async()"),
            ("agents/breach_workflow.py",        "In _research_node_async()"),
            ("agents/composed_workflow.py",      "In critical_response_node() AND standard_response_node()"),
            ("agents/supervisor_workflow.py",    "In _run_supervisor_async()"),
            ("agents/supervisor_hitl.py",        "In _phase1_async()"),
        ],
        ["File", "Location"],
    )
    doc.add_paragraph(
        "The identical loop for matching tool call inputs to ToolMessage outputs is duplicated 5+ times. "
        "Fix: move to agents/utils.py as extract_tool_calls(messages) -> list[dict]."
    )

    add_cross_cutting_table(doc,
        "4. _make_specialist_tools() and _run_in_thread() copy-pasted between two files — Major DRY",
        [
            ("_make_specialist_tools()", "supervisor_workflow.py", "supervisor_hitl.py", "Prompts slightly shortened in HITL version"),
            ("_run_in_thread()",         "supervisor_workflow.py", "supervisor_hitl.py", "Identical"),
        ],
        ["Function", "File 1", "File 2", "Difference"],
    )
    doc.add_paragraph(
        "A bug fix to any specialist agent must be applied in two places. "
        "Fix: move both to agents/specialists.py."
    )

    add_cross_cutting_table(doc,
        "5. _CVE_SERVER path computed in 6 files — DRY + DIP",
        [
            ("agents/breach_triage_agent.py",),
            ("agents/breach_workflow.py",),
            ("agents/composed_workflow.py",),
            ("agents/supervisor_workflow.py",),
            ("agents/supervisor_hitl.py",),
            ("tools/chain.py",),
        ],
        ["File"],
    )
    doc.add_paragraph(
        "Each file independently computes the absolute path to cve_server.py using os.path. "
        "Fix: add CVE_SERVER_PATH to app/config.py."
    )

    add_cross_cutting_table(doc,
        "6. MCP client configuration duplicated in 6 files — DRY + DIP",
        [
            ("agents/breach_triage_agent.py",),
            ("agents/breach_workflow.py",),
            ("agents/composed_workflow.py",),
            ("agents/supervisor_workflow.py",),
            ("agents/supervisor_hitl.py",),
            ("tools/chain.py",),
        ],
        ["File"],
    )
    doc.add_paragraph(
        "MultiServerMCPClient is constructed with the same dict in every file. "
        "Adding a second MCP server requires modifying 6 files. "
        "Fix: a factory function in agents/utils.py."
    )

    # ── SUMMARY TABLE ─────────────────────────────────────────────────────────
    doc.add_heading("Summary — Clean to Worst", level=1)

    summary_rows = [
        ("app/config.py",                       "Clean",      "—"),
        ("agents/a2a_client.py",                "Clean",      "—"),
        ("observability/fault_tolerance.py",    "Clean",      "—"),
        ("mcp_servers/cve_server.py",           "Clean",      "OCP (minor)"),
        ("agents/multimodal_agent.py",          "Minor",      "DIP"),
        ("guardrails/prompt_injection.py",      "Minor",      "DIP"),
        ("observability/logging_setup.py",      "Minor",      "SRP (minor)"),
        ("app/llm.py",                          "Minor",      "SRP (minor), DIP (minor)"),
        ("rag/retriever.py",                    "Minor",      "DIP (minor), ISP (minor)"),
        ("tools/tools.py",                      "Minor",      "OCP (minor)"),
        ("a2a_server/threat_intel_server.py",   "Moderate",   "DIP (major), SRP (moderate)"),
        ("app/rag_chain.py",                    "Moderate",   "DIP (major), SRP (moderate)"),
        ("agents/breach_triage_agent.py",       "Moderate",   "DIP (major), SRP (moderate)"),
        ("agents/breach_workflow.py",           "Moderate",   "DIP (major), SRP (moderate)"),
        ("tools/chain.py",                      "Significant","DIP (major), SRP (major)"),
        ("agents/composed_workflow.py",         "Significant","DIP (major), SRP (major), OCP (moderate)"),
        ("agents/supervisor_workflow.py",       "Significant","DIP (major), SRP (major)"),
        ("ui/app.py",                           "Significant","SRP (major), OCP (moderate)"),
        ("rag/ingestion.py",                    "Severe",     "DIP (major), SRP (major), OCP (moderate)"),
        ("agents/supervisor_hitl.py",           "Severe",     "DIP (major), SRP (major), OCP (moderate), massive duplication"),
    ]

    table = doc.add_table(rows=1 + len(summary_rows), cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["Module", "Overall", "Key Violations"])
    set_col_width(table, 0, 6.0)
    set_col_width(table, 1, 2.5)
    set_col_width(table, 2, 7.0)

    severity_to_colour = {
        "Clean":       GREEN,
        "Minor":       YELLOW,
        "Moderate":    ORANGE,
        "Significant": RGBColor(0xC5, 0x50, 0x0E),
        "Severe":      RED,
    }

    for i, (mod, overall, violations) in enumerate(summary_rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = mod
        cells[1].text = overall
        cells[2].text = violations
        for cell in cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
        colour = severity_to_colour.get(overall, YELLOW)
        set_cell_bg(cells[1], colour)
        if cells[1].paragraphs[0].runs:
            cells[1].paragraphs[0].runs[0].font.color.rgb = WHITE
            cells[1].paragraphs[0].runs[0].bold = True
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i % 2 == 0:
            set_cell_bg(cells[0], LIGHT_BLUE)
            set_cell_bg(cells[2], LIGHT_BLUE)

    doc.add_paragraph()

    # ── HIGHEST-LEVERAGE FIXES ────────────────────────────────────────────────
    doc.add_heading("Highest-Leverage Fixes (in order)", level=1)

    fixes = [
        ("1", "Centralise build_llm()",
         "Eliminates 6 copies of LLM construction. Single import change per module once the factory is right."),
        ("2", "Create agents/utils.py",
         "Add extract_text(), extract_tool_calls(), and build_mcp_client(). Eliminates the three biggest DRY violations in agent code."),
        ("3", "Add CVE_SERVER_PATH to app/config.py",
         "Eliminates the os.path.abspath(...) path computation from 6 files."),
        ("4", "Create agents/specialists.py",
         "Move _make_specialist_tools() and _run_in_thread() out of the two supervisor files that copy them."),
        ("5", "Split rag/ingestion.py",
         "Separate PDF loading, chunking strategy, vision summarisation, and ChromaDB storage into dedicated modules."),
        ("6", "Split tools/chain.py",
         "Separate MCP setup, message building, and the agentic loop into distinct concerns."),
    ]

    table = doc.add_table(rows=1 + len(fixes), cols=3)
    table.style = "Table Grid"
    add_header_row(table, ["#", "Fix", "Impact"])
    set_col_width(table, 0, 0.8)
    set_col_width(table, 1, 5.0)
    set_col_width(table, 2, 9.7)

    for i, (num, fix, impact) in enumerate(fixes, start=1):
        cells = table.rows[i].cells
        cells[0].text = num
        cells[1].text = fix
        cells[2].text = impact
        for cell in cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
        if i % 2 == 0:
            for cell in cells:
                set_cell_bg(cell, LIGHT_BLUE)

    out = "docs/solid_audit.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
